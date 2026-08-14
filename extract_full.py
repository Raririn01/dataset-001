from docx import Document

doc = Document(r'C:\Homework\LAB\dataset-001-20260619T100050Z-3-001\dataset-001\IWAIT2027_Extended_Abstract.docx')

print("=== PARAGRAPHS ===")
for p in doc.paragraphs:
    if p.text.strip():
        print(p.text)
    else:
        print()

print("\n=== TABLES ===")
for i, table in enumerate(doc.tables):
    print(f"\n--- Table {i+1} ---")
    for row in table.rows:
        cells = [cell.text for cell in row.cells]
        print(cells)

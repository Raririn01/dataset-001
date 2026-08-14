"""Generate IWAIT 2027 one-page extended abstract (.docx) from project results."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

OUTPUT = Path(__file__).parent / "IWAIT2027_Extended_Abstract.docx"

# --- Edit these before submission ---
AUTHORS = "[Your Name] and [Co-author Name if any]"
AFFILIATION = "[Your University/Department], [City], [Country] (e-mail: your.email@university.ac.th)"

TITLE = (
    "Benchmarking Pose Estimation Models for Physical "
    "Rehabilitation Exercise Assessment"
)

ABSTRACT = (
    "This study evaluates three vision-based pose estimation models—"
    "YOLOv8-pose, MoveNet SinglePose Thunder, and MediaPipe Pose Landmarker Heavy—"
    "for physical rehabilitation exercise assessment. The models were benchmarked "
    "on the UCO Physical Rehabilitation dataset using OptiTrack motion capture as "
    "ground truth. A total of 61 video-camera samples covering 13 rehabilitation "
    "exercises were processed, including lower-body and upper-body movements "
    "performed in seated, supine, and standing positions. For each video, joint "
    "landmarks were detected and rehabilitation-relevant joint angles were computed "
    "using exercise-specific joint mappings (hip-knee-ankle for lower-body and "
    "shoulder-elbow-wrist for upper-body exercises). Evaluation metrics included "
    "detection rate and mean absolute error (MAE) of estimated joint angles."
)

INTRO = (
    "Vision-based pose estimation enables low-cost monitoring of rehabilitation "
    "exercises without specialized motion-capture hardware. However, model "
    "performance varies across body regions, camera viewpoints, and patient "
    "positions. This work provides a systematic comparison of three widely used "
    "pose estimation models against OptiTrack ground truth on a multi-exercise "
    "rehabilitation dataset, offering practical guidance for selecting models in "
    "clinical rehabilitation applications."
)

METHOD = (
    "We evaluated YOLOv8-pose (Ultralytics), MoveNet SinglePose Thunder "
    "(TensorFlow), and MediaPipe Pose Landmarker Heavy on the UCO Physical "
    "Rehabilitation dataset. Each exercise was associated with fixed metadata "
    "defining body region, side, and patient position. Joint angles were computed "
    "from detected landmarks and compared frame-by-frame with OptiTrack reference "
    "angles. Supine exercises required 90° frame rotation prior to inference. "
    "MediaPipe timestamps were computed from frame indices to ensure temporal "
    "consistency. Detection rate was defined as the percentage of frames with "
    "valid angle estimates, and MAE measured angular deviation in degrees."
)

RESULTS_INTRO = (
    "Table I summarizes global performance across 61 evaluated video-camera "
    "samples spanning 13 exercise IDs (exercises 10 and 11 excluded due to "
    "missing ground truth)."
)

CONCLUSION = (
    "MediaPipe Heavy achieved the lowest angle MAE (17.46°) with a detection "
    "rate of 93.37%, closely followed by YOLOv8-pose (18.18° MAE, 93.41% "
    "detection). MoveNet Thunder showed substantially lower robustness (65.09% "
    "detection, 24.16° MAE). Both YOLOv8-pose and MediaPipe Heavy are suitable "
    "candidates for vision-based rehabilitation monitoring, while MoveNet Thunder "
    "is less effective for this multi-view rehabilitation dataset."
)

REFERENCES = [
    "[1] UCO Physical Rehabilitation Dataset, University of Córdoba.",
    "[2] J. Redmon et al., \"Ultralytics YOLOv8,\" 2023.",
    "[3] Google, \"MoveNet: Ultra fast and accurate human pose estimation,\" TensorFlow Hub.",
    "[4] Google, \"MediaPipe Pose Landmarker,\" 2023.",
    "[5] OptiTrack, \"Motion capture systems,\" NaturalPoint Inc.",
]


def set_run_font(run, size=9, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_centered(doc, text, size=11, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    return p


def add_section(doc, heading, body):
    p = doc.add_paragraph()
    run = p.add_run(heading)
    set_run_font(run, size=9, bold=True)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)

    p2 = doc.add_paragraph()
    run2 = p2.add_run(body)
    set_run_font(run2, size=9)
    p2.paragraph_format.space_after = Pt(2)
    p2.paragraph_format.first_line_indent = Inches(0.14)
    return p2


def add_table(doc):
    caption = doc.add_paragraph()
    cap_run = caption.add_run("TABLE I.  GLOBAL MODEL COMPARISON ON UCO REHABILITATION DATASET")
    set_run_font(cap_run, size=8, bold=True)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(1)

    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    headers = ["Model", "Detection Rate (%)", "Angle MAE (°)"]
    data = [
        ["YOLOv8-pose", "93.41", "18.18"],
        ["MoveNet Thunder", "65.09", "24.16"],
        ["MediaPipe Heavy", "93.37", "17.46"],
    ]
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=8, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, row in enumerate(data, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                set_run_font(run, size=8)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(9)

    add_centered(doc, TITLE, size=11, bold=True)
    add_centered(doc, AUTHORS, size=10, bold=False)
    add_centered(doc, AFFILIATION, size=9, bold=False)

    p_abs = doc.add_paragraph()
    run_abs_label = p_abs.add_run("Abstract—")
    set_run_font(run_abs_label, size=9, bold=True)
    run_abs_body = p_abs.add_run(ABSTRACT)
    set_run_font(run_abs_body, size=9)
    p_abs.paragraph_format.space_after = Pt(3)

    add_section(doc, "I. INTRODUCTION", INTRO)
    add_section(doc, "II. METHOD", METHOD)
    add_section(doc, "III. EXPERIMENTAL RESULTS", RESULTS_INTRO)
    add_table(doc)
    add_section(doc, "IV. CONCLUSION", CONCLUSION)

    p_ref = doc.add_paragraph()
    run_ref = p_ref.add_run("REFERENCES")
    set_run_font(run_ref, size=9, bold=True)
    p_ref.paragraph_format.space_before = Pt(2)

    for ref in REFERENCES:
        pr = doc.add_paragraph()
        rr = pr.add_run(ref)
        set_run_font(rr, size=8)
        pr.paragraph_format.space_after = Pt(0)
        pr.paragraph_format.left_indent = Inches(0.14)
        pr.paragraph_format.first_line_indent = Inches(-0.14)

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    build_document()
